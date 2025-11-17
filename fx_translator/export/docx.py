"""
Экспорт в DOCX формат.
"""

from __future__ import annotations
import re
from typing import Dict, Tuple, List, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from fx_translator.core.models import PageBatch
from fx_translator.utils.geometry import sort_segments_reading_order


def _sanitize_for_xml(text: str) -> str:
    """Удаляет несовместимые с XML символы."""
    if not isinstance(text, str):
        return ""

    illegal_xml_chars_re = re.compile(
        "["
        "\x00-\x08"  # C0 control codes
        "\x0b\x0c"  # Vertical tab, form feed
        "\x0e-\x1f"  # More control codes
        "\x7f-\x84"  # Delete + C1 controls
        "\x86-\x9f"  # More C1 controls
        "]"
    )
    return illegal_xml_chars_re.sub("", text)


def _soft_wrap_tokens(t: str, max_token: int = 40, insert_every: int = 20) -> str:
    """Вставляет мягкие переносы в длинные слова."""
    out_parts = []
    for tok in t.split():
        if len(tok) <= max_token:
            out_parts.append(tok)
        else:
            chunks = [
                tok[i : i + insert_every] for i in range(0, len(tok), insert_every)
            ]
            out_parts.append("\\u200b".join(chunks))
    return " ".join(out_parts)


def _font_size_for_segment(seg) -> int:
    """Определяет размер шрифта для сегмента."""
    if getattr(seg, "lineheight", 0.0) and seg.lineheight > 0:
        return int(max(8.0, min(18.0, float(seg.lineheight))))

    seg_type = (seg.type or "").lower()
    if "title" in seg_type:
        return 16
    if "section" in seg_type or "header" in seg_type:
        return 14
    if "caption" in seg_type:
        return 9
    if "footnote" in seg_type:
        return 8
    if "list" in seg_type:
        return 11
    return 12


def export_docx(
    pages: List[PageBatch],
    translations: Dict[Tuple[int, str, int], str],
    out_docx: str,
    title: Optional[str] = None,
) -> None:
    """Экспортирует страницы с переводами в DOCX."""
    doc = Document()

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(11)

    if title:
        p = doc.add_paragraph(_sanitize_for_xml(title))
        p.style = doc.styles["Title"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for page_batch in pages:
        side = getattr(page_batch, "logicalside", "")
        side_suffix = f"[{side}]" if side in ("L", "R") else ""

        h = doc.add_paragraph(f"Страница {page_batch.pagenumber}{side_suffix}")
        h.style = doc.styles["Heading 1"]

        table = doc.add_table(rows=1, cols=4)
        table.allow_autofit = False

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Тип"
        hdr_cells[2].text = "Оригинал"
        hdr_cells[3].text = "Перевод"

        col_widths = [Inches(0.6), Inches(1.0), Inches(3.5), Inches(3.5)]
        for col, w in zip(table.columns, col_widths):
            for cell in col.cells:
                cell.width = w

        for s in sort_segments_reading_order(page_batch.segments):
            row = table.add_row().cells
            row[0].text = str(s.blockid)
            row[1].text = _sanitize_for_xml(s.type)

            orig_text = _sanitize_for_xml(s.text)
            orig_text_wrapped = _soft_wrap_tokens(orig_text)
            row[2].paragraphs[0].add_run(orig_text_wrapped)

            tr_text = translations.get((page_batch.pagenumber, side, s.blockid), "")
            tr_text_sanitized = _sanitize_for_xml(tr_text)
            tr_text_wrapped = _soft_wrap_tokens(tr_text_sanitized)
            run = row[3].paragraphs[0].add_run(tr_text_wrapped)
            run.font.size = Pt(_font_size_for_segment(s))

    try:
        doc.save(out_docx)
    except Exception as e:
        import logging

        logging.error(f"Не удалось сохранить DOCX файл {out_docx}: {e}")
        raise
